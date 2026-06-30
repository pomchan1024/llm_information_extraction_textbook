! pip install python-docx==1.2.0

import docx

def extract_text_from_word(file_path):
    doc = docx.Document(file_path)
    full_text = []

    # すべての段落からテキストを取得
    for para in doc.paragraphs:
        if para.text.strip(): # 空行を除外する場合
            full_text.append(para.text)

    # すべての表からテキストを取得
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return "\n".join(full_text)

# --- 実行例 ---
text = extract_text_from_word("sample_data/sample.docx")
print(text)